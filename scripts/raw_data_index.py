import os
from docx import Document
import re

def create_numbered_version(input_path):
    # 创建新文件名
    base_name = os.path.basename(input_path)
    new_name = base_name.replace(".docx", "_numbered.docx")
    output_path = os.path.join(os.path.dirname(input_path), new_name)

    doc = Document(input_path)
    numbered_doc = Document()  # 新建空白文档
    
    current_number = 1
    number_pattern = re.compile(r'^\[\d+\]')
    
    # 处理段落（修复样式复制问题）
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 跳过空行
        if not text:
            numbered_doc.add_paragraph()
            continue
            
        # 检查已有编号
        if number_pattern.match(text):
            new_text = text
        else:
            new_text = f'[{current_number}] {text}'
            current_number += 1
        
        # 创建新段落并复制格式
        new_para = numbered_doc.add_paragraph()
        new_run = new_para.add_run(new_text)
        
        # 仅当原段落有格式时复制
        if para.runs:
            source_run = para.runs[0]
            new_run.bold = source_run.bold
            new_run.italic = source_run.italic
            if source_run.font.name:
                new_run.font.name = source_run.font.name
            if source_run.font.size:
                new_run.font.size = source_run.font.size

    numbered_doc.save(output_path)
    return output_path

if __name__ == "__main__":
    input_file = "/Volumes/Study/prj/data/raw/training_raw_data.docx"
    output_file = create_numbered_version(input_file)
    print(f"已创建带编号的新文件：{output_file}")