import os
import re
import time
import random
from typing import List, Dict
from docx import Document
from docx.shared import Pt
from collections import defaultdict
import dashscope
from dashscope import Generation

# 配置信息
CONFIG = {
    "api_key": "sk-3b986ed51abb4ed18aadde5d41e11397",
    "base_url": "https://dashscope-intl.aliyuncs.com/compatible-mode/v1",
    "input_path": "/Volumes/Study/prj/data/processed/filtered_results.docx",
    "output_path": "/Volumes/Study/prj/data/raw/training_labeled_data.docx",
    "label_rules": {
        "高频": ["明亮", "齿音", "空气感", "干净"],
        "中频": ["人声厚度", "鼻音", "浑浊感", "饱和感"],
        "低频": ["Bass平衡", "厚重感", "低频"],
        "压缩": ["动态控制", "句头", "句尾"],
        "reverb": ["空间感", "环境效果", "混响"],
        "声场": ["宽度", "定位", "立体感"],
        "音量": ["电平调整", "音量"],
        "效果器": ["电", "autotune", "电话音", "失真"]
    },
    "conflict_rules": [
        (["高频", "低频"], 0.3),
        (["压缩", "reverb"], 0.5)
        
    ],
    "verify_per_label": 1  # 每个标签验证样本数
}

class AudioPreLabelSystem:
    def __init__(self):
        dashscope.api_key = CONFIG["api_key"]
        dashscope.base_url = CONFIG["base_url"]
        self.prelabeler = self.AudioPreLabeler(CONFIG["label_rules"], CONFIG["conflict_rules"])
        self.samples = []
        self.verified_data = []
        self.label_stats = defaultdict(int)
        self.start_time = time.time()

    class AudioPreLabeler:
        def __init__(self, label_rules: Dict[str, List[str]], conflict_rules: List[tuple]):
            self.label_rules = label_rules
            self.conflict_rules = conflict_rules
            self._build_index()

        def _build_index(self):
            """构建多级索引体系"""
            self.keyword_index = defaultdict(list)  # 精确匹配索引
            self.pattern_index = []                # 正则模式索引
            
            # 第一层：精确关键词匹配
            for label, keywords in self.label_rules.items():
                for kw in keywords:
                    self.keyword_index[kw].append(label)
            
            # 第二层：模糊匹配规则
            self._add_fuzzy_rules()
            
            # 第三层：上下文关联规则
            self._add_context_rules()
        
        def _add_fuzzy_rules(self):
            """添加模糊匹配规则"""
            fuzzy_patterns = [
                (r'太亮|刺耳', '高频'),
                (r'闷|不清晰', '中频'),
                (r'轰|震', '低频'),
                (r'压(缩|得太)|动态', '压缩'),
                (r'空间|大厅', 'reverb'),
                (r'声场|立体', '声场'),
                (r'音量|太(大|小)', '音量'),
                (r'甜|圆润', '高频'),
                (r'磁性|厚', '中频'),
                (r'颓废|教堂', 'reverb'),
                (r'贴耳|凸出', '声场'),
                (r'layback|后置', 'reverb'),
                (r'干净|电', '高频')
            ]
            for pattern, label in fuzzy_patterns:
                self.pattern_index.append((re.compile(pattern), label))
        
        def _add_context_rules(self):
            """添加上下文关联规则"""
            self.context_rules = [
                (r'人声.*伴奏', ['声场', '音量']),
                (r'低频.*气势', ['低频', '压缩']),
                (r'Autotune', ['效果器']),
                (r'punchy', ['中频', '压缩']),
                (r'空间感', ['reverb', '声场'])
            ]
        
        def _resolve_conflicts(self, labels: List[str]) -> List[str]:
            """解决标签冲突"""
            valid_labels = labels.copy()
            for conflict_pair, penalty in self.conflict_rules:
                if all(label in valid_labels for label in conflict_pair):
                    valid_labels = [l for l in valid_labels if l not in conflict_pair[1:]]
            return valid_labels
        
        def prelabel(self, text: str) -> List[str]:
            """执行多级标注流程"""
            # 第一阶段：精确匹配
            exact_matches = []
            for kw, labels in self.keyword_index.items():
                if kw in text:
                    exact_matches.extend(labels)
            
            # 第二阶段：模糊匹配
            pattern_matches = []
            for regex, label in self.pattern_index:
                if regex.search(text):
                    pattern_matches.append(label)
            
            # 第三阶段：上下文增强
            context_matches = []
            for pattern, labels in self.context_rules:
                if re.search(pattern, text):
                    context_matches.extend(labels)
            
            # 合并结果并去重
            all_labels = list(set(exact_matches + pattern_matches + context_matches))
            
            # 冲突解决
            final_labels = self._resolve_conflicts(all_labels)
            
            return final_labels[:3]   # 最多返回3个最相关标签
    
    def _call_qwen(self, text: str) -> List[str]:
        """适配QwQ模型的流式调用方法"""
        prompt = f"""作为专业混音师，请将用户需求转换为以下标准标签：
    可选标签包括：{', '.join(CONFIG['label_rules'].keys())}
    输入描述：{text}
    输出格式：仅返回逗号分隔的标签列表，不要其他内容"""
        
        try:
            # 修改后的QwQ调用
            response = Generation.call(
                model="qwq-plus-2025-03-05",  # 模型名称变更
                messages=[{
                    "role": "user",
                    "content": prompt
                }],
                stream=True,                # 启用流式
                incremental_output=True,    # 必须参数
                temperature=0.2,
                top_p=0.7,
                max_tokens=50
            )
            
            # 流式响应处理
            full_output = ""
            for chunk in response:
                if chunk.status_code == 200:
                    full_output += chunk.output.choices[0]['message']['content']
                else:
                    print(f"流式块错误: {chunk.code}")
            
            # 后续处理保持原样
            if not full_output:
                print("API返回空内容")
                return []
                
            return [l.strip() for l in full_output.split(",") 
                    if l.strip() in CONFIG["label_rules"]][:3]
            
        except Exception as e:
            print(f"API请求失败: {str(e)}")
            return []

    def _load_input(self) -> List[str]:
        """加载输入文件"""
        try:
            if not os.path.exists(CONFIG["input_path"]):
                raise FileNotFoundError(f"文件未找到: {CONFIG['input_path']}")
                 
            doc = Document(CONFIG["input_path"])
            raw_texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            print(f"成功加载 {len(raw_texts)} 条有效数据")
            return raw_texts
        except Exception as e:
            print(f"文件加载失败: {str(e)}")
            exit(1)

    def _save_output(self):
        """保存输出文件"""
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        
        # 创建紧凑列表格式
        line_num = 1
        for item in self.verified_data:
            # 将标签列表转换为逗号分隔的中文字符串
            labels_str = "，".join(item['labels'])
            p = doc.add_paragraph(f"{line_num}. {item['text']}")
            p.add_run(f"（{labels_str}）").bold = True  # 将所有标签放在同一个括号内
            line_num += 1  # 每个样本递增行号
        
        # 保存文件
        os.makedirs(os.path.dirname(CONFIG["output_path"]), exist_ok=True)
        doc.save(CONFIG["output_path"])
        print(f"\n已保存预标注结果到：{CONFIG['output_path']}")
    def manual_verification(self):
        """New verification process"""
        label_map = defaultdict(list)
        for sample in self.samples:
            for label in sample['labels']:
                label_map[label].append(sample)
        
        # First round verification: Sample 1 from each category
        verify_samples = []
        for label in CONFIG["label_rules"].keys():
            samples = label_map.get(label, [])
            if samples:
                verify_samples.append(random.choice(samples))
        
        print(f"\n{'='*50}")
        print(f"Starting label sampling verification (total {len(verify_samples)} samples)")
        
        for idx, sample in enumerate(verify_samples, 1):
            print(f"\n[Verification Progress {idx}/{len(verify_samples)}]")
            print(f"Original Description: {sample['text']}")
            print(f"Current Labels: {', '.join(sample['labels'])}")
            action = input("Accept? (y/n/e): ").strip().lower()
            
            if action == 'e':
                self._edit_labels(sample)
            elif action == 'y':
                self.verified_data.append(sample)
        
        # 自动接受未验证样本
        auto_accept = [s for s in self.samples 
                       if not any(s['text'] == v['text'] for v in self.verified_data)]
        self.verified_data.extend(auto_accept)
        
        self._show_statistics()

    def _show_statistics(self):
        """显示统计信息"""
        self.label_stats = defaultdict(int)
        for sample in self.verified_data:
            for label in sample['labels']:
                self.label_stats[label] += 1
        
        print("\n" + "="*50)
        print(f"总耗时：{time.time()-self.start_time:.2f}秒")
        print(f"总样本数：{len(self.verified_data)}")
        for label, count in self.label_stats.items():
            print(f"{label}：{count}条")

    def _edit_labels(self, sample: dict):
        """编辑标签"""
        new_labels = []
        print(f"当前标签： {sample['labels']}")
        print("可用标签：", ", ".join(self.prelabeler.label_rules.keys()))
        
        while True:
            input_labels = input("请输入正确标签（空格分隔）: ").split()
            valid_labels = [l for l in input_labels if l in self.prelabeler.label_rules]
            
            if valid_labels:
                sample['labels'] = valid_labels[:3]  # 更新标签并限制最多3个
                self.verified_data.append(sample)
                print(f"✓ 已更新为：{valid_labels}")
                break
            print("包含无效标签，请重新输入！")

if __name__ == "__main__":
    system = AudioPreLabelSystem()
    
    try:
        # 加载数据
        raw_texts = system._load_input()
        
        # 预标注阶段
        for idx, text in enumerate(raw_texts, 1):
            labels = system.prelabeler.prelabel(text)
            if not labels:
                labels = system._call_qwen(text)
            system.samples.append({
                'text': text,
                'labels': labels[:3]
            })
            print(f"\r标注进度: {idx}/{len(raw_texts)}", end="", flush=True)
        print("\n预标注完成")
        
        # 验证阶段
        system.manual_verification()
        
        # 最终保存
        system._save_output()
        system._show_statistics()
        
    except KeyboardInterrupt:
        print("\n用户中断操作，正在保存已验证数据...")
        system._save_output()
        system._show_statistics()
    except Exception as e:
        print(f"系统错误: {str(e)}")
        exit(1)