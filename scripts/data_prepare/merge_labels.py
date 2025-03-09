import re
import os
from docx import Document
from collections import OrderedDict
import sys

def process_labeled_data(input_path):
    try:
        # ===== 1. 验证文件路径 =====
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"文件不存在：{input_path}")
        if not input_path.endswith('.docx'):
            raise ValueError("只支持 .docx 文件格式")
        if not os.access(input_path, os.R_OK):
            raise PermissionError(f"文件不可读：{input_path}")
        if not os.access(input_path, os.W_OK):
            print("警告：没有原文件的写入权限，尝试复制到临时文件处理...")

        # ===== 2. 读取文档内容 =====
        doc = Document(input_path)
        pattern = re.compile(r"\d+\.\s+(.*?)\s*（(.*?)）")
        merged_data = OrderedDict()

        # ===== 3. 合并数据 =====
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            match = pattern.match(text)
            if match:
                content = match.group(1).strip()
                tag = match.group(2).strip()
                
                if content in merged_data:
                    merged_data[content].add(tag)
                else:
                    merged_data[content] = {tag}

        # ===== 4. 重建文档结构 =====
        new_doc = Document()
        
        # 保留原文档格式（标题、样式等）
        if doc.paragraphs:
            first_para = doc.paragraphs[0]
            new_doc.add_paragraph(first_para.text, style=first_para.style.name)
        
        # 添加合并后的内容
        for idx, (content, tags) in enumerate(merged_data.items(), 1):
            sorted_tags = "，".join(sorted(tags))
            new_text = f"{idx}. {content}（{sorted_tags}）"
            new_doc.add_paragraph(new_text)

        # ===== 5. 保存文件 =====
        try:
            new_doc.save(input_path)
            print(f"成功更新文件：{os.path.abspath(input_path)}")
        except PermissionError:
            temp_path = os.path.join(os.path.dirname(input_path), "TEMP_" + os.path.basename(input_path))
            new_doc.save(temp_path)
            print(f"因权限问题保存到临时文件：{temp_path}")
            print("请手动替换原文件")

    except Exception as e:
        print(f"处理失败：{str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    # ===== 路径配置 =====
    current_script = os.path.abspath(__file__)  # 当前脚本绝对路径
    prj_root = os.path.dirname(os.path.dirname(os.path.dirname(current_script)))  # 回退三级到 /Volumes/Study/prj
    
    input_path = os.path.join(
        prj_root,
        "data",
        "raw",
        "training_labeled_data.docx"
    )
    
    print("="*40)
    print(f"项目根目录：{prj_root}")
    print(f"目标文件路径：{input_path}")
    print("="*40 + "\n")

    process_labeled_data(input_path)