import os
import math
import json
import random
from collections import defaultdict
from docx import Document
import logging

# 配置日志
logging.basicConfig(
    filename='entropy_optimizer.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

class EntropyOptimizer:
    def __init__(self, input_json, output_dir):
        self.input_json = input_json
        self.output_dir = output_dir
        self.data = []
        self.original_entropy = 0.0
        self.original_dist = {}
        
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 加载清洗后的数据
        try:
            with open(input_json, 'r', encoding='utf-8') as f:
                self.data = json.load(f)
            logging.info(f"成功加载 {len(self.data)} 条清洗后数据")
            print(f"加载 {len(self.data)} 条清洗后数据")
        except Exception as e:
            logging.error(f"加载数据失败: {str(e)}")
            raise ValueError("无法加载输入 JSON 文件，请检查路径或文件格式。")

    def calculate_entropy(self, data):
        """修正后的加权熵计算（考虑多标签样本的影响）"""
        label_weights = defaultdict(float)
        total_weight = 0.0
        
        for item in data:
            # 使用标签数量的倒数作为权重（多标签样本权重更低）
            weight = 1.0 / len(item['labels'])
            for label in item['labels']:
                label_weights[label] += weight
            total_weight += weight
        
        entropy = 0.0
        for weight in label_weights.values():
            prob = weight / total_weight
            if prob > 0:
                entropy -= prob * math.log2(prob)
        
        return entropy, label_weights



    def interactive_optimization(self):
        """改进的交互式优化"""
        current_entropy, _ = self.calculate_entropy(self.data)
        self.original_entropy = current_entropy
        self.original_dist = self.get_label_distribution()
        
        print(f"当前数据熵值: {current_entropy:.4f}")
        target = input(f"请输入目标熵值 (当前值={current_entropy:.4f}, 建议范围2.3-2.7，直接回车跳过优化): ")
        
        if not target.strip():
            print("跳过优化流程")
            return self.data
        
        try:
            target_entropy = float(target)
            max_possible = math.log2(len(self.original_dist))
            if target_entropy >= max_possible:
                print(f"目标熵值不可高于理论最大值 {max_possible:.4f}")
                return self.data
        except ValueError:
            print("无效输入，跳过优化")
            return self.data
        
        return self.optimize(target_entropy)

    def optimize(self, target_entropy):
        """改进的优化逻辑（避免熵值异常上升）"""
        iteration = 0
        max_iter = 50
        prev_entropy = math.inf
        tolerance = 0.005  # 允许的误差范围
        prev_dist = None
        
        while iteration < max_iter:
            current_entropy, label_dist = self.calculate_entropy(self.data)
            
            # 早停检查：熵值不再下降或出现异常波动
            if (abs(current_entropy - target_entropy) < tolerance) or \
            (prev_entropy - current_entropy < tolerance) or \
            (current_entropy > prev_entropy):
                print(f"达到目标熵值或熵值不再下降，停止优化")
                break
            
            prev_entropy = current_entropy
            prev_dist = label_dist.copy()
            
            # 选择高频标签（基于原始标签数量，非权重）
            label_counts = self.get_label_distribution()
            avg_count = len(self.data) / len(label_counts)
            candidates = [label for label, count in label_counts.items() 
                        if count > avg_count * 1.5]
            
            if not candidates:
                candidates = sorted(label_counts.items(), key=lambda x: -x[1])[:2]
            
            # 优先删除高频标签的单标签样本
            for label in candidates[:2]:
                items = [item for item in self.data 
                        if label in item['labels'] and len(item['labels']) == 1]
                
                if len(items) <= 3:
                    continue
                
                remove_num = max(1, int(len(items)*0.3))  # 提高删除比例
                self.data = [item for item in self.data 
                            if item not in random.sample(items, remove_num)]
                
                print(f"迭代{iteration+1}: 删除{remove_num}条单标签'{label}'样本")
            
            iteration += 1
        
        # 若优化后熵值反而上升，回滚到上一次状态
        final_entropy, _ = self.calculate_entropy(self.data)
        if final_entropy > prev_entropy:
            self.data = [item for item in self.data 
                        if item in prev_dist]
            print("检测到熵值异常上升，已自动回滚到上一次优化状态")
        
        return self.data

    def save_results(self, optimized_data):
        """改进的保存方法"""
        # 保存JSON
        output_json = os.path.join(self.output_dir, 'optimized_data.json')
        try:
            with open(output_json, 'w', encoding='utf-8') as f:
                json.dump(optimized_data, f, indent=2, ensure_ascii=False)
            logging.info(f"优化数据已保存为 JSON: {output_json}")
        except Exception as e:
            logging.error(f"保存 JSON 文件失败: {str(e)}")
            raise
        
        # 保存DOCX
        doc = Document()
        for item in optimized_data:
            p = doc.add_paragraph(f"{item['id']}. {item['text']}")
            p.add_run(f"（{'、'.join(item['labels'])}）").bold = True
        
        docx_path = os.path.join(self.output_dir, 'optimized_data.docx')
        try:
            doc.save(docx_path)
            logging.info(f"优化数据已保存为 DOCX: {docx_path}")
            print(f"优化数据已保存至: {self.output_dir}")
        except Exception as e:
            logging.error(f"保存 DOCX 文件失败: {str(e)}")
            raise

    def generate_comparison_report(self):
        """生成对比分析报告"""
        new_entropy, new_dist = self.calculate_entropy(self.data)
        report = {
            "original": {
                "entropy": self.original_entropy,
                "distribution": self.original_dist
            },
            "optimized": {
                "entropy": new_entropy,
                "distribution": self.get_label_distribution()
            },
            "changes": {
                "entropy_reduction": self.original_entropy - new_entropy,
                "removed_samples": len(self.original_dist) - len(self.data)
            }
        }
        
        report_path = os.path.join(self.output_dir, 'entropy_comparison.json')
        try:
            with open(report_path, 'w', encoding='utf-8') as f:
                json.dump(report, f, indent=2, ensure_ascii=False)
            logging.info(f"对比分析报告已保存: {report_path}")
        except Exception as e:
            logging.error(f"保存对比分析报告失败: {str(e)}")
            raise

    def get_label_distribution(self):
        """获取当前标签分布"""
        dist = defaultdict(int)
        for item in self.data:
            for label in item['labels']:
                dist[label] += 1
        return dict(dist)

if __name__ == "__main__":
    input_json = '/Volumes/Study/prj/data/processed/cleaned_data.json'
    output_dir = '/Volumes/Study/prj/data/optimized'
    
    try:
        optimizer = EntropyOptimizer(input_json, output_dir)
        optimized_data = optimizer.interactive_optimization()
        optimizer.save_results(optimized_data)
        optimizer.generate_comparison_report()
    except Exception as e:
        logging.critical(f"程序运行失败: {str(e)}")
        print(f"程序运行失败: {str(e)}")