import os
import re
from docx import Document

def get_script_dir():
    return os.path.dirname(os.path.abspath(__file__))

def clean_text(text):
    text = re.sub(r"[-—_]{10,}", "", text)
    text = re.sub(r'[→─_]+', '', text).strip()
    return text

def read_backtrans(file_path):
    doc = Document(file_path)
    buffer = [clean_text(para.text.strip()) for para in doc.paragraphs if para.text.strip()]
    
    data = []
    for i in range(0, len(buffer), 6):
        group = buffer[i:i+6]
        if len(group) == 6:
            entry = {}
            # 使用更强大的正则表达式模式
            entry["1"] = re.sub(r'^【原始】\s*', '', group[0]).strip()
            entry["2"] = re.sub(r'^\d+\.\s*英译[\s→]*', '', group[1]).strip()
            entry["3"] = re.sub(r'^\d+\.\s*回译[\s→]*', '', group[2]).strip()
            entry["4"] = re.sub(r'^\d+\.\s*专业[\s→]*', '', group[3]).strip()
            entry["5"] = re.sub(r'^\d+\.\s*通俗[\s→]*', '', group[4]).strip()
            entry["6"] = re.sub(r'^\d+\.\s*反义[\s→]*', '', group[5]).strip()
            data.append(entry)
    
    log_info = {
        'file': os.path.basename(file_path),
        'entries': len(data),
        'paragraphs': len(buffer)
    }
    return data, log_info

def read_augmented(file_path):
    doc = Document(file_path)
    paragraphs = [clean_text(para.text.strip()) for para in doc.paragraphs if para.text.strip()]
    
    data = []
    i = 0
    while i < len(paragraphs):
        if paragraphs[i].startswith("原始数据："):
            entry = {"1": paragraphs[i].replace("原始数据：", "").strip()}
            i += 1
            current_key = 2
            
            while i < len(paragraphs) and re.match(r'^EDA变体\d+：', paragraphs[i]):
                entry[f"{current_key}"] = re.sub(r'^EDA变体\d+：', '', paragraphs[i]).strip()
                current_key += 1
                i += 1
            
            data.append(entry)
        else:
            i += 1
    
    log_info = {
        'file': os.path.basename(file_path),
        'entries': len(data),
        'paragraphs': len(paragraphs)
    }
    return data, log_info

def merge_and_save(data1, data2, output_path):
    doc = Document()
    global_id = 1
    
    # 合并回译数据
    for entry in data1:
        for key in sorted(entry.keys(), key=lambda x: int(x)):
            text = entry[key].strip()
            if text and not text.startswith(('1.', '2.', '3.', '4.', '5.')):
                doc.add_paragraph(f"{global_id}. {text}")
                global_id += 1
    
    # 合并增强数据
    for entry in data2:
        for key in sorted(entry.keys(), key=lambda x: int(x)):
            text = entry[key].strip()
            if text:
                doc.add_paragraph(f"{global_id}. {text}")
                global_id += 1
    
    doc.save(output_path)
    return global_id - 1

if __name__ == "__main__":
    script_dir = get_script_dir()
    input_files = {
        'backtrans': "../../data/processed/backtrans_results.docx",
        'augmented': "../../data/processed/augmented_results.docx"
    }
    output_path = os.path.join(script_dir, "../../data/processed/merged_results.docx")
    
    data1, log1 = read_backtrans(os.path.join(script_dir, input_files['backtrans']))
    data2, log2 = read_augmented(os.path.join(script_dir, input_files['augmented']))
    
    total_paragraphs = merge_and_save(data1, data2, output_path)
    
    print("====== 处理完成 ======")
    print(f"回译文件处理：{log1['entries']} 条目，{log1['paragraphs']} 段落")
    print(f"增强文件处理：{log2['entries']} 条目，{log2['paragraphs']} 段落")
    print(f"合并文档已保存至：{os.path.abspath(output_path)}")
    print(f"总生成段落数：{total_paragraphs}")