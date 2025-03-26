"""音乐需求增强处理脚本 v3.0
(支持五种数据增强方式)

论文核心引用：
1. 受限采样（Restricted Sampling）：
   "Restricted sampling appears to be the most consistent approach, always scaling with larger monolingual data." (Section 5.3)
2. 禁用标签平滑：
   "Disabling label smoothing for the target-to-source model [...] results in higher-quality synthetic data." (Section 4.1)
3. N-best列表采样：
   "50-best sampling improves significantly in both test sets." (Section 5.3)
"""
"""音乐需求增强处理脚本 v3.1（稳定五结果版）"""
"""音乐需求增强处理脚本 v3.2（论文优化版）"""

import os
import re
import sys
from pathlib import Path
import time
from datetime import datetime
from docx import Document
from docx.shared import Pt
import dashscope
from dashscope import Generation

# ======================
# 基础配置（论文参数优化）
# ======================
BASE_DIR = Path("/Volumes/Study/prj")
CONFIG = {
    "api_key": "sk-3b986ed51abb4ed18aadde5d41e11397",
    "input_path": BASE_DIR / "data/raw/training_raw_data.docx",
    "output_path": BASE_DIR / "data/processed/backtrans_results.docx",
    "log_path": BASE_DIR / "data/processed/processing.log",
    
    # 论文参数优化
    "sampling": {
        "base_temp": 0.3,     # 基础温度（对应禁用标签平滑）
        "high_temp": 0.6,     # 反义生成温度（N-best采样）
        "top_p": 0.85,        # 受限采样阈值
        "max_retry": 3        # 最大重试次数
    },
    
    # 正则过滤模式（置信度掩码）
    "patterns": {
        "antonym": r"(相反的描述：|可以这样表述：|这意味着).*",
        "synonym": r"(改写后的句子：|建议改为：)",
        "default": r"[\(（].*?[\)）]"
    }
}

class RobustAugmenter:
    def __init__(self):
        # 初始化API连接
        dashscope.api_key = CONFIG["api_key"]
        dashscope.base_url = "https://dashscope-intl.aliyuncs.com/compatible-mode/v1"
        
        # 文件验证
        if not CONFIG["input_path"].exists():
            raise FileNotFoundError(f"输入文件不存在：{CONFIG['input_path']}")
        CONFIG["output_path"].parent.mkdir(parents=True, exist_ok=True)
        
        # 增强统计系统
        self.stats = {
            "total": 0, "success": 0, "failed": 0,
            "input_tokens": 0, "output_tokens": 0,
            "retries": 0
        }
        self.start_time = time.time()

    def _call_api(self, prompt: str, text: str, is_critical: bool = False) -> str:
        """API calling (with restart)"""
        params = {
            "model": "qwen-max",
            "prompt": f"{prompt}\n{text}",
            "temperature": CONFIG["sampling"]["high_temp"] if is_critical else CONFIG["sampling"]["base_temp"],
            "top_p": CONFIG["sampling"]["top_p"],
            "max_tokens": 100  
        }
        
        for attempt in range(CONFIG["sampling"]["max_retry"]):
            try:
                response = Generation.call(**params)
                
                if not hasattr(response, 'output') or not response.output.text:
                    raise ValueError("无效API响应结构")
                
                if hasattr(response, 'usage'):
                    self.stats["input_tokens"] += response.usage.input_tokens
                    self.stats["output_tokens"] += response.usage.output_tokens
                
                return response.output.text.strip()
                
            except Exception as e:
                error_type = type(e).__name__
                print(f"API异常（{attempt+1}/{CONFIG['sampling']['max_retry']}）[{error_type}]: {str(e)}")
                self.stats["retries"] += 1
                time.sleep(2 ** attempt)
        
        return "[生成失败]"

    def _clean_output(self, text: str, mode: str) -> str:
        """后处理模块（基于论文置信度过滤）"""
        # 模式匹配冗余内容
        patterns = {
            "antonym": CONFIG["patterns"]["antonym"],
            "synonym": CONFIG["patterns"]["synonym"],
            "default": CONFIG["patterns"]["default"]
        }
        
        # 分步清洗
        clean_text = re.sub(patterns.get(mode, patterns["default"]), "", text)
        clean_text = re.sub(r"\s+", " ", clean_text)  # 合并多余空格
        clean_text = clean_text.strip("”“\"'：、")  # 去除边缘符号
        
        # 截断处理（保持核心语义）
        sentences = re.split(r'[。！？]', clean_text)
        return sentences[0].strip() if sentences else clean_text

    def _generate_results(self, text: str) -> tuple:
        """生成优化后的五结果（带论文方法约束）"""
        # 反向翻译增强（Section 3.1）
        en_trans = self._clean_output(
            self._call_api("严格翻译为英文（仅输出结果）：", text),
            "default"
        )
        back_trans = self._clean_output(
            self._call_api("将英文回译为中文（仅输出结果）：", en_trans),
            "default"
        )
        
        # 同义替换（Section 4.2.1）
        synonym_pro = self._clean_output(
            self._call_api("用专业音乐术语改写（保持原意）：", text),
            "synonym"
        )
        synonym_pop = self._clean_output(
            self._call_api("用通俗语言简化表达：", text),
            "synonym"
        )
        
        # 反义生成（Section 4.2.2）
        antonym = self._clean_output(
            self._call_api("生成完全相反的音乐需求（仅输出需求本身）：", text, is_critical=True),
            "antonym"
        )
        
        return (en_trans, back_trans, synonym_pro, synonym_pop, antonym)

    def _save_report(self):
        """增强版报告生成"""
        elapsed = time.time() - self.start_time
        report = f"""
{' 音乐需求增强报告 '.center(40, '=')}
处理时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
========================================
成功处理: {self.stats['success']}/{self.stats['total']} 
API重试: {self.stats['retries']}次
Token用量: 输入 {self.stats['input_tokens']} | 输出 {self.stats['output_tokens']}
平均耗时: {elapsed/(self.stats['success'] or 1):.2f}秒/条
总耗时: {self._format_duration(elapsed)}
========================================"""
        with open(CONFIG["log_path"], "w") as f:
            f.write(report)
        print(report)

    def _format_duration(self, seconds: float) -> str:
        """格式化时间差"""
        hours, rem = divmod(seconds, 3600)
        minutes, seconds = divmod(rem, 60)
        return f"{int(hours):02}:{int(minutes):02}:{int(seconds):02}"

    def process(self):
        """主处理流程（优化版）"""
        print("🎵 启动音乐需求增强（论文优化版）...")
        input_doc = Document(CONFIG["input_path"])
        output_doc = Document()
        
        try:
            valid_paras = [p.text.strip() for p in input_doc.paragraphs if p.text.strip()]
            total = len(valid_paras)
            
            for idx, text in enumerate(valid_paras, 1):
                self.stats["total"] += 1
                try:
                    results = self._generate_results(text)
                    
                    # 结构化写入
                    output_doc.add_paragraph(f"【原始】{text}")
                    output_doc.add_paragraph(f"1. 英译 → {results[0]}")
                    output_doc.add_paragraph(f"2. 回译 → {results[1]}")
                    output_doc.add_paragraph(f"3. 专业 → {results[2]}")
                    output_doc.add_paragraph(f"4. 通俗 → {results[3]}")
                    output_doc.add_paragraph(f"5. 反义 → {results[4]}")
                    
                    self.stats["success"] += 1
                    print(f"进度: {idx}/{total} ({idx/total:.0%}) | 最新: {results[1][:20]}...", flush=True)
                    
                    # 动态限流（1-3秒随机延迟）
                    time.sleep(min(max(idx % 3, 1), 3))
                    
                except Exception as e:
                    print(f"处理失败: {text[:15]}... | 错误: {type(e).__name__}-{str(e)}")
                    self.stats["failed"] += 1
                    
            output_doc.save(CONFIG["output_path"])
            print(f"\n✅ 结果已保存至: {CONFIG['output_path']}")
            
        except KeyboardInterrupt:
            print("\n⚠️ 用户中断！正在保存已处理内容...")
            output_doc.save(CONFIG["output_path"])
        finally:
            self._save_report()

if __name__ == "__main__":
    try:
        RobustAugmenter().process()
    except Exception as e:
        print(f"❌ 系统初始化失败: {type(e).__name__}-{str(e)}")
