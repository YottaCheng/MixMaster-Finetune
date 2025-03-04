"""音乐需求回译处理脚本 v2.4
(合并回译与反义生成)"""
import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Pt
import dashscope
from dashscope import Generation

# 配置信息
CONFIG = {
    "api_key": "sk-3b986ed51abb4ed18aadde5d41e11397",
    "base_url": "https://dashscope-intl.aliyuncs.com/compatible-mode/v1",
    "input_path": "/Volumes/Study/prj/data/raw/training_raw_data.docx",
    "output_dir": "/Volumes/Study/prj/data/processed",
    "output_file": "backtrans_results.docx",
    "log_file": "processing.log"
}

class BackTranslator:
    def __init__(self):
        # 初始化 DashScope
        dashscope.api_key = os.getenv("DASHSCOPE_API_KEY", CONFIG["api_key"])
        dashscope.base_url = CONFIG["base_url"]
        
        # 初始化统计信息
        self.usage_data = {
            "processed_items": 0,    # 实际处理的数据条数
            "total_requests": 0,     # API调用次数
            "input_tokens": 0,
            "output_tokens": 0
        }
        
        # 时间记录
        self.start_time = time.time()
        
        # 创建输出目录
        os.makedirs(CONFIG["output_dir"], exist_ok=True)
        self.output_path = os.path.join(CONFIG["output_dir"], CONFIG["output_file"])
        self.log_path = os.path.join(CONFIG["output_dir"], CONFIG["log_file"])

    def _call_qwen(self, prompt: str, text: str) -> str:
        """调用 Qwen 大模型并记录统计"""
        full_prompt = f"{prompt}\n原句：{text}"
        
        try:
            response = Generation.call(
                model="qwen-max",
                prompt=full_prompt,
                temperature=0.3,
                top_p=0.8,
                max_tokens=500
            )
            
            # 更新统计信息
            if hasattr(response, 'usage'):
                self.usage_data["total_requests"] += 1
                self.usage_data["input_tokens"] += response.usage.input_tokens
                self.usage_data["output_tokens"] += response.usage.output_tokens
            
            return response.output.text.strip()
        except Exception as e:
            print(f"API 调用失败: {str(e)}")
            return ""

    def _process_translation_and_reverse(self, text: str) -> tuple:
        """执行回译和反义生成，返回 (英译, 回译, 反义)"""
        prompt = f"""你是一个混音小白，假设你要发出请求给混音师，请完成以下任务：
1. 将以下混音调整需求翻译为英文，并再翻译回中文，尽可能保持原意和细节。
2. 将以下混音调整为相反意思。

格式要求：
英译：<翻译结果>
回译：<回译结果>
反义：<调整结果>

原句：{text}"""
        
        result = self._call_qwen(prompt, text)
        
        # 解析结果
        en_trans = ""
        back_trans = text
        reversed_text = text
        
        try:
            if "英译：" in result and "回译：" in result and "反义：" in result:
                parts = result.split("反义：")
                reversed_text = parts[-1].strip()
                
                trans_parts = parts[0].split("回译：")
                back_trans = trans_parts[-1].strip()
                
                en_part = trans_parts[0].split("英译：")[-1].strip()
                en_trans = en_part
        except Exception as e:
            print(f"解析失败: {str(e)}")
        
        return (en_trans, back_trans, reversed_text)

    def _format_duration(self, seconds: float) -> str:
        """格式化时间差为易读格式"""
        hours, remainder = divmod(seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        return f"{int(hours)}小时{int(minutes)}分{int(seconds)}秒"

    def _save_report(self):
        """生成并保存综合统计报告"""
        total_tokens = self.usage_data["input_tokens"] + self.usage_data["output_tokens"]
        total_time = self._format_duration(time.time() - self.start_time)
        
        report = f"""
处理时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
────────────────────────────────
原始数据总量：{self.usage_data["processed_items"]} 条
API请求次数：{self.usage_data["total_requests"]} 次
输入Token消耗：{self.usage_data["input_tokens"]}
输出Token消耗：{self.usage_data["output_tokens"]}
总Token消耗：{total_tokens}
处理总耗时：{total_time}
────────────────────────────────"""
        
        # 写入日志文件
        with open(self.log_path, "w", encoding="utf-8") as f:
            f.write(report)
            
        # 控制台输出
        print(report)

    def process_document(self):
        """处理整个文档"""
        print("🔄 开始处理文档...")
        input_doc = Document(CONFIG["input_path"])
        output_doc = Document()
        output_doc.styles['Normal'].font.name = '微软雅黑'
        output_doc.styles['Normal'].font.size = Pt(11)
        try:
            for idx, para in enumerate(input_doc.paragraphs, 1):
                original = para.text.strip()
                if not original:
                    continue
                
                # 记录处理进度
                self.usage_data["processed_items"] += 1
                
                # 获取所有数据
                en_trans, back_trans, reversed_text = self._process_translation_and_reverse(original)
                
                # 写入文档
                output_doc.add_paragraph(f"原始数据：{original}")
                output_doc.add_paragraph(f"回译结果：{back_trans}")
                output_doc.add_paragraph(f"英译结果：{en_trans}")
                output_doc.add_paragraph(f"反义结果：{reversed_text}\n")
                
                print(f"处理进度：第 {idx} 条")
                time.sleep(1.5)  # 避免 API 调用频率过高
            
            # 保存结果并生成报告
            output_doc.save(self.output_path)
            self._save_report()
            print(f"\n✅ 处理完成！结果已保存至：{self.output_path}")
        except Exception as e:
            print(f"\n❌ 处理中断: {str(e)}")
            self._save_report()  # 保存已处理的数据

if __name__ == "__main__":
    BackTranslator().process_document()