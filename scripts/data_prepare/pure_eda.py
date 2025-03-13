"""
音乐需求增强系统 v3.8
（路径优化版 + 增强错误处理）
"""
"""
论文参考:
1. 《A Survey of Data Augmentation Approaches for NLP》
2. 《AEDA: An Easier Data Augmentation Technique for Text Classification》
"""
import os
import json
import random
import warnings
import jieba
from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.shared import Pt

# 配置相对路径
BASE_DIR = Path(__file__).parent.parent.parent  # 假设脚本在 scripts/ 目录
CONFIG_DIR = BASE_DIR / "config"
DATA_DIR = BASE_DIR / "data"

CONFIG_PATH = CONFIG_DIR / "music_synonyms.json"
INPUT_PATH = DATA_DIR / "raw" / "training_raw_data.docx"
OUTPUT_PATH = DATA_DIR / "processed" / "augmented_results.docx"
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

class MusicAugmentor:
    def __init__(self):
        # 自动创建必要目录
        OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
        
        if not CONFIG_PATH.exists():
            raise FileNotFoundError(f"术语库文件不存在于：{CONFIG_PATH}")
        
        self.term_map = self._load_config()
        jieba.initialize()

        # 专业术语保护列表
        self.protected_terms = {
            "压缩": "[COMP]", "混响": "[REV]", "EQ": "[EQ]",
            "Autotune": "[AT]", "齿音": "[DEESS]", "底鼓": "[KICK]",
            "声场": "[SF]", "动态": "[DYN]", "频段": "[BAND]",
            "人声": "[VOCAL]", "伴奏": "[BACKING]"
        }
    
    def _load_config(self):
        """安全加载层级化术语库"""
        try:
            with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
                raw_data = json.load(f)
            
            term_map = defaultdict(list)
            for category in raw_data.values():
                for key, values in category.items():
                    if not isinstance(values, list):
                        values = [values]
                    clean_values = list(set(values))
                    # 主索引
                    term_map[key].extend(clean_values)
                    # 创建小写索引
                    term_map[key.lower()].extend(clean_values)
                    # 反向索引
                    for val in clean_values:
                        term_map[val].append(key)
            return term_map
        except json.JSONDecodeError:
            raise ValueError(f"配置文件格式错误：{CONFIG_PATH}")
        except Exception as e:
            raise RuntimeError(f"术语库加载失败：{str(e)}")
    
    def _dynamic_replace(self, text: str, seed: int = None) -> tuple:
        """智能术语替换引擎（返回替换文本和替换次数）"""
        if seed is not None:
            random.seed(seed)
        
        replace_count = 0
        words = jieba.lcut(text)
        cursor = 0
        replaced = []
        
        while cursor < len(words):
            found = False
            # 优先匹配长短语（3→2→1）
            for length in [3, 2, 1]:
                if cursor + length > len(words):
                    continue
                
                phrase = "".join(words[cursor:cursor+length])
                if replacements := self.term_map.get(phrase):
                    shuffled_replacements = random.sample(replacements, len(replacements))
                    replaced.append(shuffled_replacements[0])
                    cursor += length
                    replace_count += 1
                    found = True
                    break
            if not found:
                current = words[cursor]
                if random.random() < 0.9:
                    if replacements := self.term_map.get(current):
                        shuffled_replacements = random.sample(replacements, len(replacements))
                        replaced.append(shuffled_replacements[0])
                        replace_count += 1
                    else:
                        replaced.append(current)
                else:
                    replaced.append(current)
                cursor += 1
        return "".join(replaced), replace_count
    
    def process(self, text: str) -> dict:
        """生成增强结果（自动过滤低质量变体）"""
        # 使用不同的随机种子生成两个变体
        eda1_text, count1 = self._dynamic_replace(text, seed=random.randint(1, 1000000))
        eda2_text, count2 = self._dynamic_replace(text, seed=random.randint(1, 1000000))
        
        return {
            "original": text,
            "eda1": eda1_text if count1 >= 3 else "",
            "eda2": eda2_text if count2 >= 3 else ""
        }

def main():
    # 初始化增强器
    try:
        augmentor = MusicAugmentor()
    except Exception as e:
        print(f"初始化失败: {str(e)}")
        return
    
    # 创建输出文档
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal'].font.size = Pt(11)
    
    try:
        if not INPUT_PATH.exists():
            raise FileNotFoundError(f"输入文件不存在：{INPUT_PATH}")
            
        input_doc = Document(INPUT_PATH)
        total = len(input_doc.paragraphs)
        
        for idx, para in enumerate(input_doc.paragraphs, 1):
            if text := para.text.strip():
                result = augmentor.process(text)
                doc.add_paragraph(f"原始数据：{result['original']}")
                if result['eda1']:
                    doc.add_paragraph(f"EDA变体1：{result['eda1']}")
                if result['eda2']:
                    doc.add_paragraph(f"EDA变体2：{result['eda2']}")
                doc.add_paragraph()
                print(f"处理进度: {idx}/{total} ({idx/total:.1%})")
                
        doc.save(OUTPUT_PATH)
        print(f"\n✅ 处理完成！结果保存至: {OUTPUT_PATH}")
    except Exception as e:
        print(f"\n❌ 处理中断: {str(e)}")

if __name__ == "__main__":
    main()