"""
基于SIGMOD '21论文《Rotom: A Meta-Learned Data Augmentation Framework》的数据过滤脚本
实现功能：去重 + 置信度感知的轻度清洗（相对路径版本）
"""

import os
import re
from docx import Document
from collections import defaultdict

def get_script_dir():
    """获取脚本所在目录"""
    return os.path.dirname(os.path.abspath(__file__))

def load_document(file_path):
    """加载Word文档并提取带序号的段落"""
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 解析序号和内容（支持"1.内容"或"1内容"格式）
            match = re.match(r'^(\d+)\.?\s*(.*)', text)
            if match:
                num, content = match.groups()
                paragraphs.append((int(num), content))
    return paragraphs

def filter_duplicates(paragraphs):
    """基于内容的去重（保留第一个出现的条目）"""
    seen = set()
    unique_paragraphs = []
    dup_count = 0
    
    for num, content in paragraphs:
        if content not in seen:
            seen.add(content)
            unique_paragraphs.append((num, content))
        else:
            dup_count += 1
            
    return unique_paragraphs, dup_count

def clean_content(content):
    """
    实施轻度清洗（参考论文3.2节InvDA和4.1节过滤策略）
    1. 清理特殊符号（保留基本标点）
    2. 移除前后空白
    3. 合并连续空格
    """
    # 清理非常用符号（保留中文、英文、数字及常见标点）
    cleaned = re.sub(r'[^\w\u4e00-\u9fa5\.,!?;:\-\s\(\)]', '', content)
    # 合并连续空白字符
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned

def apply_cleaning_rules(paragraphs):
    """应用清洗规则并过滤无效条目"""
    valid_paragraphs = []
    invalid_count = 0
    
    for num, content in paragraphs:
        # 轻度清洗
        cleaned = clean_content(content)
        
        # 过滤规则（参考论文Table 3的span级操作）
        # 1. 长度过短（<5字符）
        # 2. 无效内容（纯符号或数字）
        if len(cleaned) < 5:
            invalid_count += 1
            continue
            
        if re.match(r'^[\d\s\W]+$', cleaned):  # 无有效文字
            invalid_count += 1
            continue
            
        valid_paragraphs.append((num, cleaned))
        
    return valid_paragraphs, invalid_count

def save_filtered_doc(paragraphs, output_path):
    """保存处理后的文档并重新编号"""
    doc = Document()
    for idx, (orig_num, content) in enumerate(paragraphs, 1):
        doc.add_paragraph(f"{idx}. {content}")
    doc.save(output_path)

def main(input_path, output_path):
    # 加载数据
    raw_paragraphs = load_document(input_path)
    print(f"[1/4] 载入数据完成 | 原始条目数: {len(raw_paragraphs)}")
    
    # 去重处理
    unique_paras, dup_count = filter_duplicates(raw_paragraphs)
    print(f"[2/4] 去重完成 | 删除重复条目: {dup_count} | 剩余条目: {len(unique_paras)}")
    
    # 实施清洗
    cleaned_paras, invalid_count = apply_cleaning_rules(unique_paras)
    print(f"[3/4] 清洗完成 | 过滤无效条目: {invalid_count} | 剩余条目: {len(cleaned_paras)}")
    
    # 保存结果
    save_filtered_doc(cleaned_paras, output_path)
    print(f"[4/4] 结果已保存至: {os.path.abspath(output_path)}")
    
    # 生成统计报告
    stats = {
        "original": len(raw_paragraphs),
        "duplicates": dup_count,
        "invalid": invalid_count,
        "final": len(cleaned_paras)
    }
    print("\n=== 统计报告 ===")
    print(f"原始条目: {stats['original']}")
    print(f"删除重复: {stats['duplicates']}")
    print(f"过滤无效: {stats['invalid']}")
    print(f"最终保留: {stats['final']} ({stats['final']/stats['original']:.1%})")

if __name__ == "__main__":

    script_dir = get_script_dir()
    base_dir = os.path.abspath(os.path.join(script_dir, "../../"))  # 项目根目录
    
    input_file = os.path.join(base_dir, "data/processed/merged_results.docx")
    output_file = os.path.join(base_dir, "data/processed/filtered_results.docx")
    
    main(input_file, output_file)