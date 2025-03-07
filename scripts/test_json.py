import json
import os
import re
from docx import Document

def extract_label(text):
    """从文本中提取括号内的标签"""
    pattern = r'^(.*?)\s*（([^）]*)）$'
    match = re.match(pattern, text)
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return text.strip(), None

def docx_to_jsonl(input_path, output_path):
    doc = Document(input_path)
    records = []
    
    # 处理段落
    for para in doc.paragraphs:
        full_text = para.text.strip()
        if not full_text:
            continue
            
        input_text, label = extract_label(full_text)
        if label:
            records.append({
                "instruction": "请对以下文本进行分类",
                "input": input_text,
                "output": label
            })
    
    # 处理表格（保持原有逻辑）
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(cells) >= 2:
                records.append({
                    "instruction": cells[0],
                    "input": cells[1],
                    "output": cells[2] if len(cells) > 2 else "待补充"
                })
    
    # 保存文件
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        for record in records:
            f.write(json.dumps(record, ensure_ascii=False) + '\n')
    
    print(f"转换完成！共生成 {len(records)} 条有效数据")
    print(f"输出路径：{output_path}")

if __name__ == "__main__":
    input_docx = r"D:\kings\MixMaster-Finetune\data\raw\test_raw_data.docx"
    output_json = r"D:\kings\MixMaster-Finetune\LLaMA-Factory\data\mix\validation_data.jsonl"
    docx_to_jsonl(input_docx, output_json)