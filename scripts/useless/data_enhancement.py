# data_enhancement.py
import os
import json
import time
import shutil
import random
import docx
import torch
from tqdm import tqdm
from pathlib import Path
from transformers import AutoTokenizer, AutoModelForCausalLM
import logging

# ----------------------
# 配置模块
# ----------------------
class AugmentConfig:
    """动态配置管理"""
    def __init__(self):
        # 路径配置
        self.project_root = Path("/Volumes/Study/prj")
        self.raw_path = self.project_root / "data/raw/training_raw_data.docx"
        self.aug_path = self.project_root / "data/raw/augmented_data.docx"
        self.backup_dir = self.project_root / "data/backups"
        self.synonym_path = self.project_root / "config/synonyms.json"
        
        # 模型配置
        self.model_name = "deepseek-ai/DeepSeek-R1-Distill-Qwen-1.5B"  # 改为较小模型
        self.num_aug = 3
        self.max_length = 60
        self.temperature = 0.8
        self.top_p = 0.9
        
        # 设备配置
        self.device = "mps" if torch.backends.mps.is_available() else "cpu"
        
        # 初始化环境
        self._setup_env()
    
    def _setup_env(self):
        """初始化目录和配置文件"""
        self.backup_dir.mkdir(parents=True, exist_ok=True)
        if not self.synonym_path.exists():
            self._init_default_synonyms()
    
    def _init_default_synonyms(self):
        """生成默认同义词库"""
        default = {
            "人声": ["嗓音", "歌声", "演唱", "声线"],
            "混响": ["残响", "空间效果", "环境音"],
            "贴耳": ["邻近感", "亲密感", "零距离"]
        }
        with open(self.synonym_path, 'w', encoding='utf-8') as f:
            json.dump(default, f, ensure_ascii=False, indent=2)

# ----------------------
# 数据处理模块
# ----------------------
class DataProcessor:
    @staticmethod
    def read_docx(path):
        """安全读取文档内容"""
        try:
            doc = docx.Document(path)
            return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        except Exception as e:
            logging.error(f"文件读取失败: {path} | 错误: {str(e)}")
            return []
    
    @staticmethod
    def save_docx(texts, path, config):
        """带版本控制的保存"""
        try:
            # 创建带时间戳的备份
            timestamp = int(time.time())
            backup_path = config.backup_dir / f"aug_backup_{timestamp}.docx"
            if Path(path).exists():
                shutil.copy(path, backup_path)
            
            # 保存新数据
            doc = docx.Document()
            seen = set()
            for text in texts:
                if text and text not in seen:
                    doc.add_paragraph(text)
                    seen.add(text)
            doc.save(path)
            logging.info(f"已保存 {len(seen)} 条数据到 {path}")
            
            # 清理旧备份（保留最近5个）
            backups = sorted(config.backup_dir.glob("aug_backup_*.docx"), key=os.path.getmtime)
            for old_backup in backups[:-5]:
                old_backup.unlink()
        except Exception as e:
            logging.error(f"文件保存失败: {str(e)}")

# ----------------------
# 增强引擎
# ----------------------
# data_enhancement.py
class AugmentationEngine:
    def __init__(self, config):
        self.config = config
        self.synonyms = self.load_synonyms()
        self.model, self.tokenizer = self.load_model()

    def load_model(self):
        """优化后的模型加载方法"""
        try:
            # 显式指定设备
            device = torch.device(self.config.device)
            
            tokenizer = AutoTokenizer.from_pretrained(self.config.model_name)
            model = AutoModelForCausalLM.from_pretrained(
                self.config.model_name,
                torch_dtype=torch.float32,  # 使用float32兼容MPS
                device_map={"": device},    # 显式指定设备
                trust_remote_code=True
            )
            
            # MPS内存优化配置
            if device.type == "mps":
                torch.mps.set_per_process_memory_fraction(0.8)  # 限制内存使用
                torch.mps.empty_cache()
            
            logging.info(f"成功加载模型到 {device}")
            return model.to(device), tokenizer
        except Exception as e:
            logging.error(f"模型加载失败: {str(e)}")
            logging.error("建议解决方案：")
            logging.error("1. 升级PyTorch到最新版本: pip install --pre torch torchvision torchaudio")
            logging.error("2. 使用较小模型: 修改config.model_name为'deepseek-ai/DeepSeek-R1-Distill-Qwen-1B'")
            raise

    # ...其余代码保持不变...
    def load_synonyms(self):
        """动态加载同义词库"""
        try:
            with open(self.config.synonym_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logging.error(f"同义词库加载失败: {str(e)}")
            return {}
    
    def dynamic_replace(self, text):
        """智能同义词替换"""
        words = list(text)
        for i in range(len(words)):
            if random.random() < 0.6:  # 60%替换概率
                replacements = self.synonyms.get(words[i], [])
                if replacements:
                    words[i] = random.choice(replacements)
        return ''.join(words)
    
    def generate_variants(self, text):
        """生成增强变体"""
        prompt = f"[音乐制作需求改写] 原始: {text}\n改写:"
        try:
            inputs = self.tokenizer(prompt, return_tensors="pt").to(self.model.device)
            outputs = self.model.generate(
                inputs.input_ids,
                max_new_tokens=self.config.max_length,
                num_return_sequences=self.config.num_aug,
                temperature=self.config.temperature,
                top_p=self.config.top_p,
                do_sample=True,
                pad_token_id=self.tokenizer.eos_token_id
            )
            generations = [
                self.tokenizer.decode(output, skip_special_tokens=True)
                .split("改写:")[-1]
                .strip()
                for output in outputs
            ]
            return list(set([
                self.dynamic_replace(text) 
                for text in generations 
                if 8 <= len(text) <= 20
            ]))
        except Exception as e:
            logging.error(f"生成失败: {text[:15]}... | 错误: {str(e)[:50]}")
            return []

# ----------------------
# 主流程
# ----------------------
def main():
    # 初始化配置
    config = AugmentConfig()
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
        handlers=[logging.StreamHandler()]
    )
    
    # 读取数据
    raw_data = DataProcessor.read_docx(config.raw_path)
    if not raw_data:
        logging.error("输入文件为空或不存在！")
        return
    
    # 初始化引擎
    try:
        engine = AugmentationEngine(config)
    except:
        return
    
    # 执行增强
    enhanced = []
    with tqdm(raw_data, desc="处理进度", unit="条") as pbar:
        for text in pbar:
            variants = engine.generate_variants(text)
            if variants:
                enhanced.extend(variants)
                enhanced.append(text)  # 保留原始文本
                pbar.set_postfix({"最新生成": variants[0][:15] + "..."})
    
    # 保存结果
    DataProcessor.save_docx(list(set(enhanced)), config.aug_path, config)

if __name__ == "__main__":
    main()