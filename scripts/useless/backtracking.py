from googletrans import Translator
from docx import Document

# 初始化翻译器
translator = Translator()

def back_translate(text, src_lang='zh-cn'):
    """
    使用回译方法对文本进行数据增强。
    1. 中文 -> 英文
    2. 英文 -> 中文
    """
    try:
        # 第一步：中文 -> 英文
        en_translation = translator.translate(text, src=src_lang, dest='en').text
        
        # 第二步：英文 -> 中文
        zh_translation = translator.translate(en_translation, src='en', dest='zh-cn').text
        
        return zh_translation
    except Exception as e:
        print(f"翻译失败: {e}")
        return None

# 读取原始数据文件
def read_raw_data(file_path):
    """
    从 .docx 文件中读取原始数据。
    :param file_path: 原始数据文件路径
    :return: 包含每行数据的列表
    """
    doc = Document(file_path)
    sentences = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return sentences

# 保存增强后的数据到新文件
def save_augmented_data(sentences, output_file):
    """
    将增强后的句子保存到新的 .docx 文件中。
    :param sentences: 增强后的句子列表
    :param output_file: 输出文件路径
    """
    doc = Document()
    for sentence in sentences:
        doc.add_paragraph(sentence)
    doc.save(output_file)

# 主函数
if __name__ == "__main__":
    # 输入文件路径
    input_file = "/Volumes/Study/prj/data/raw/training_raw_data.docx"
    # 输出文件路径
    output_file = "/Volumes/Study/prj/data/raw/training_data_augmented.docx"

    # 读取原始数据
    original_sentences = read_raw_data(input_file)

    # 数据增强
    augmented_data = []
    for sentence in original_sentences:
        augmented_sentence = back_translate(sentence)
        if augmented_sentence:
            augmented_data.append(augmented_sentence)

    # 保存增强后的数据
    save_augmented_data(augmented_data, output_file)

    print(f"处理完成！已将增强后的数据保存到：{output_file}")