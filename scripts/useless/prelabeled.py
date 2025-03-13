import os
import re
import time
import random
import numpy as np
import requests
import os
from dashscope import api_key
from typing import List, Dict
from docx import Document
from docx.shared import Pt
from collections import defaultdict
import dashscope
from dashscope import Generation
import pandas as pd 
from concurrent.futures import ThreadPoolExecutor, as_completed
from tenacity import retry, stop_after_attempt, wait_exponential
from snorkel.labeling import PandasLFApplier
from snorkel.labeling import LabelingFunction
from snorkel.labeling.model import LabelModel

# 配置信息（已更新API信息）
CONFIG = {
    "api_key": "sk-3511f72cb3324a36b42ac8dc91568769",  
    "base_url":"https://dashscope.aliyuncs.com/compatible-mode/v1",
    "model": "deepseek-r1",
    "input_path": "/Volumes/Study/prj/data/processed/filtered_results.docx",
    "output_path": "/Volumes/Study/prj/data/raw/training_labeled_data.docx",
    "auto_save_interval": 50,
    "label_rules": {
        "高频": ["明亮", "齿音", "空气感", "干净"],
        "中频": ["人声厚度", "鼻音", "浑浊感", "饱和感"],
        "低频": ["Bass平衡", "厚重感", "低频"],
        "压缩": ["动态控制", "句头", "句尾"],
        "reverb": ["空间感", "环境效果", "混响"],
        "声场": ["宽度", "定位", "立体感"],
        "音量": ["电平调整", "音量"],
        "效果器": ["autotune","电话音","失真"]
    },
    "conflict_rules": [
        (["高频", "低频"], 0.3),
        (["压缩", "reverb"], 0.5)
    ],
    "gen_model_params": {
        "epochs": 100,
        "lr": 0.01,
        "metric": "f1"
    },
    "verify_per_label": 1
}

class AudioPreLabelSystem:
    @retry(stop=stop_after_attempt(3), 
          wait=wait_exponential(multiplier=1, min=2, max=10))
    def __init__(self):
        dashscope.api_key = CONFIG["api_key"]
        dashscope.base_url = CONFIG["base_url"]  # 确保base_url正确
        
        # 修改测试请求
        try:
            test_response = Generation.call(
                model=CONFIG["model"],  # 使用配置中的模型名称
                messages=[{"role": "user", "content": "test"}],
                result_format='message',  # 添加必要参数
                timeout=5
            )
            if test_response.status_code == 200:
                print("✅ API连接测试通过")
            else:
                raise Exception(f"测试失败: {test_response.message}")
        except Exception as e:
            print(f"❌ API连接失败: {str(e)}")
            exit(1)
        
        print("🔄 初始化标注函数...")
        self.lfs = self._init_labeling_functions()
        print(f"✅ 已加载 {len(self.lfs)} 个标注函数")
        
        print("🔄 初始化生成模型...")
        self.label_model = LabelModel(cardinality=len(CONFIG["label_rules"]))
        print("✅ 生成模型初始化完成")
        
        self.L_train = None
        self.samples = []
        self.verified_data = []
        self.label_stats = defaultdict(int)
        self.start_time = time.time()
        self.applier = PandasLFApplier(lfs=self.lfs)


    def _init_labeling_functions(self):
        """修正闭包问题和API函数位置"""
        lfs = []
        
        # 修正关键字规则闭包
        for label, keywords in CONFIG["label_rules"].items():
            def make_kw_lf(l=label, kws=keywords):  # ✅ 固化变量
                def lf_func(x):
                    return l if any(kw in x.text for kw in kws) else -1
                return LabelingFunction(name=f"kw_{l}", f=lf_func)
            lfs.append(make_kw_lf())
        
        # 转换正则规则
        regex_rules = [
            (r'太亮|刺耳', '高频'),
            (r'闷|不清晰', '中频'),
            (r'轰|震', '低频')
        ]
        for pattern, label in regex_rules:
            def make_re_lf(p=pattern, l=label):   # 绑定当前值
                def lf_func(x):
                    return l if re.search(p, x.text) else -1
                return LabelingFunction(name=f"re_{l}", f=lf_func)
            lfs.append(make_re_lf())
        
        
        
        # 添加API弱监督
        def api_lf(x):
            labels = self._call_deepseek(x.text)
            return labels[0] if labels else -1
        lfs.append(LabelingFunction(name="api_weak", f=api_lf))
        
        return lfs

    class TextWrapper:
        """适配原有文本格式到Snorkel"""
        def __init__(self, text: str):
            self.text = text

    def _train_generative_model(self, texts: List[str]):
        from tqdm import tqdm
        
        print(f"\n🔄 开始生成弱监督标签（共{len(texts)}条）")
        with tqdm(total=len(texts), desc="生成进度") as pbar:
            with ThreadPoolExecutor(max_workers=5) as executor:
                futures = {executor.submit(self.applier.apply, pd.DataFrame({"text": [text]})): text for text in texts}
                results = []
                for future in as_completed(futures):
                    results.extend(future.result())
                    pbar.update(1)
                self.L_train = np.array(results)
        for i, future in enumerate(as_completed(futures)):
            results.extend(future.result())
            pbar.update(1)
            
            # 新增自动保存逻辑
            if (i+1) % CONFIG["auto_save_interval"] == 0:
                self._save_checkpoint(results, i+1)

    def _custom_analysis_report(self):
        """自定义标注函数质量分析"""
        print("\n=== 标注函数质量分析 ===")
        coverage = (self.L_train != -1).mean(axis=0)
        conflicts = (self.L_train[:, None] != self.L_train) & (self.L_train != -1)
        
        print(f"{'函数名称':<20} | {'覆盖率':<10} | {'冲突率':<10}")
        print("-" * 50)
        for i, lf in enumerate(self.lfs):
            conflict_rate = conflicts[:, i].mean()
            print(f"{lf.name:<20} | {coverage[i]:<10.2%} | {conflict_rate:<10.2%}")

    def _call_deepseek(self, text: str) -> List[str]:
        """修正后的API调用方法"""
        prompt = f"""作为专业混音师，请将以下需求转换为标准标签：
可选标签：{", ".join(CONFIG['label_rules'].keys())}
输入描述：{text}
只需返回逗号分隔的标签列表，不要其他内容"""
        
        try:
            response = Generation.call(
                model=CONFIG["model"],  # 使用配置中的模型名称
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                top_p=0.7,
                max_tokens=50,
                result_format='message',  # 必须参数
                timeout=15
            )
            
            # 修正响应解析逻辑
            if response.status_code == 200:
                content = response.output.choices[0].message.content
                return [
                    label.strip()
                    for label in content.split(",")
                    if label.strip() in CONFIG["label_rules"]
                ][:3]
            else:
                print(f"API错误: {response.code} - {response.message}")
                return []
                

        except Exception as e:
            print(f"‼️ API请求失败详情：")
            print(f"URL: {dashscope.base_url}")
            print(f"Model: deepseek-chat")
            print(f"Error: {str(e)}")
            print(f"请求内容: {text[:50]}...")  # 显示部分输入内容
            return []
        
    def _load_input(self) -> List[str]:
        """保持原有数据加载逻辑"""
        try:
            doc = Document(CONFIG["input_path"])
            return [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        except Exception as e:
            print(f"文件加载失败: {str(e)}")
            exit(1)

    def _save_output(self, probs: np.ndarray = None):
        """增强保存逻辑（保持原有格式）"""
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)

        line_num = 1
        for idx, item in enumerate(self.verified_data):
            text = item['text']
            labels = item['labels']
            
            # 添加概率信息（如果存在）
            prob_info = ""
            if probs is not None:
                top_probs = sorted(zip(CONFIG["label_rules"].keys(), probs[idx]), 
                               key=lambda x: x[1], reverse=True)[:3]
                prob_info = " | ".join([f"{k}:{v:.2f}" for k,v in top_probs])
            
            p = doc.add_paragraph(f"{line_num}. {text}")
            p.add_run(f"（{', '.join(labels)} {prob_info}）").bold = True
            line_num += 1
        
        os.makedirs(os.path.dirname(CONFIG["output_path"]), exist_ok=True)
        doc.save(CONFIG["output_path"])
        print(f"\n已保存预标注结果到：{CONFIG['output_path']}")

    def manual_verification(self):
        """全新验证流程"""
        label_map = defaultdict(list)
        for sample in self.samples:
            for label in sample['labels']:
                label_map[label].append(sample)
        
        # 首轮验证：每类抽1个样本
        verify_samples = []
        for label in CONFIG["label_rules"].keys():
            samples = label_map.get(label, [])
            if samples:
                verify_samples.append(random.choice(samples))
        
        print(f"\n{'='*50}")
        print(f"开始标签抽样验证（共{len(verify_samples)}个样本）")
        
        for idx, sample in enumerate(verify_samples, 1):
            print(f"\n【验证进度 {idx}/{len(verify_samples)}】")
            print(f"原始描述：{sample['text']}")
            print(f"当前标签：{', '.join(sample['labels'])}")
            action = input("是否接受？(y/n/e): ").strip().lower()
            
            if action == 'e':
                self._edit_labels(sample)
            elif action == 'y':
                self.verified_data.append(sample)
        
        # 自动接受未验证样本
        auto_accept = [s for s in self.samples 
                       if not any(s['text'] == v['text'] for v in self.verified_data)]
        self.verified_data.extend(auto_accept)
        
        self._show_statistics()
    def _save_checkpoint(self, results, processed_num):
        checkpoint_path = os.path.join(os.path.dirname(CONFIG["output_path"]), "checkpoint.pkl")
        with open(checkpoint_path, 'wb') as f:
            pickle.dump({
                'processed_num': processed_num,
                'results': results,
                'timestamp': time.time()
            }, f)
        print(f"\n🔔 已保存检查点（已处理{processed_num}条）")
        
    def run(self):
        """增强后的主流程"""
        try:
            # 阶段1：数据加载
            raw_texts = self._load_input()
            
            # 阶段2：训练生成模型
            self._train_generative_model(raw_texts)
            
            # 阶段3：生成概率标签
            probs = self._generate_probs()
            
            # 阶段4：生成初始标注
            self.samples = [{
                'text': text,
                'labels': [list(CONFIG["label_rules"].keys())[np.argmax(p)]],
                'prob': max(p)
            } for text, p in zip(raw_texts, probs)]
            
            # 阶段5：人工验证（保持原有逻辑）
            self.manual_verification()
            
            # 阶段6：保存结果（增强输出）
            self._save_output(probs)
            self._show_statistics()

        except KeyboardInterrupt:
            print("\n用户中断操作，正在保存已验证数据...")
            self._save_output()
            self._show_statistics()



    def _show_statistics(self):
        """显示统计信息"""
        self.label_stats = defaultdict(int)
        for sample in self.verified_data:
            for label in sample['labels']:
                self.label_stats[label] += 1
        
        print("\n" + "="*50)
        print(f"总耗时：{time.time()-self.start_time:.2f}秒")
        print(f"总样本数：{len(self.verified_data)}")
        for label, count in self.label_stats.items():
            print(f"{label}：{count}条")

    def _edit_labels(self, sample: dict):
        """编辑标签"""
        new_labels = []
        print(f"当前标签： {sample['labels']}")
        #print("可用标签：", ", ".join(self.prelabeler.label_rules.keys()))
        print("可用标签：", ", ".join(CONFIG["label_rules"].keys()))  # ✅ 直接使用CONFIG
        valid_labels = [l for l in input_labels if l in CONFIG["label_rules"]]
        
        while True:
            input_labels = input("请输入正确标签（空格分隔）: ").split()
            valid_labels = [l for l in input_labels if l in self.prelabeler.label_rules]
            
            if valid_labels:
                sample['labels'] = valid_labels[:3]  # 更新标签并限制最多3个
                self.verified_data.append(sample)
                print(f"✓ 已更新为：{valid_labels}")
                break
            print("包含无效标签，请重新输入！")
    def _generate_probs(self):  # 添加缺失方法
        """生成概率标签"""
        return self.label_model.predict_proba(self.L_train)   

if __name__ == "__main__":
    print("🚀 启动音频预标注系统")
    system = AudioPreLabelSystem()
    
    try:
        print("\n=== 阶段1：数据加载 ===")
        raw_texts = system._load_input()
        print(f"📥 已加载 {len(raw_texts)} 条原始数据")

        print("\n=== 阶段2：生成模型训练 ===")
        system._train_generative_model(raw_texts)

        print("\n=== 阶段3：概率标签生成 ===")
        probs = system._generate_probs()
        print(f"🔖 已生成 {probs.shape[0]} 条概率标签")

        print("\n=== 阶段4：初始标注生成 ===")
        system.samples = [{
            'text': text,
            'labels': [list(CONFIG["label_rules"].keys())[np.argmax(p)]],
            'prob': max(p)
        } for text, p in zip(raw_texts, probs)]
        print("🖍️ 初始标注完成")

        print("\n=== 阶段5：人工验证 ===")
        system.manual_verification()

        print("\n=== 阶段6：结果保存 ===")
        system._save_output(probs)
        system._show_statistics()
        print("🎉 处理流程完成！")

    except KeyboardInterrupt:
        print("\n⚠️ 用户中断操作，正在保存已验证数据...")
        system._save_output()
        system._show_statistics()