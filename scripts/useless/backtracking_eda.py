"""
音乐需求增强系统 v3.4
（精准语义版）
"""
import os
import json
import random
import warnings
import jieba
import torch
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from transformers import pipeline
from functools import lru_cache

# 配置常量
CONFIG_PATH = "/Volumes/Study/prj/config/music_synonyms.json"
INPUT_PATH = "/Volumes/Study/prj/data/raw/training_raw_data.docx"
OUTPUT_PATH = "/Volumes/Study/prj/data/processed/augmented_results.docx"

# 抑制警告
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

class MusicAugmentor:
    def __init__(self):
        if not os.path.exists(CONFIG_PATH):
            raise FileNotFoundError(f"术语库文件不存在于：{CONFIG_PATH}")
        
        self.translation_cache = {}  # 增加翻译缓存
        self.term_map = self._load_config()
        self.device = self._get_device()
        jieba.initialize()

        # 初始化翻译管道
        self.translator = pipeline(
            "translation_zh_to_en",
            model="Helsinki-NLP/opus-mt-zh-en",
            device=self.device,
            num_beams=5,                
            max_length=400,
            truncation=True
        )
        self.back_translator = pipeline(
            "translation_en_to_zh",
            model="Helsinki-NLP/opus-mt-en-zh",
            device=self.device,
            num_beams=6
        )
        
        # 专业术语保护列表
        self.protected_terms = {
            "压缩": "[COMP]", "混响": "[REV]", "EQ": "[EQ]",
            "Autotune": "[AT]", "齿音": "[DEESS]", "底鼓": "[KICK]",
            "声场": "[SF]", "动态": "[DYN]", "频段": "[BAND]",
            "人声": "[VOCAL]", "伴奏": "[BACKING]", "Autotune": "[AT]",
            "压缩": "[COMP]", "混响": "[REV]", "声场": "[SF]"
        }

    def _get_device(self):
        """设备选择优化"""
        if torch.cuda.is_available():
            return "cuda"
        elif torch.backends.mps.is_available():
            torch.mps.empty_cache()
            return "mps"
        return "cpu"

    def _load_config(self):
        """层级化术语库加载"""
        try:
            with open(CONFIG_PATH) as f:
                raw_data = json.load(f)
            
            term_map = defaultdict(list)
            for category in raw_data.values():
                for key, values in category.items():
                    clean_values = list(set(values)) if isinstance(values, list) else [values]
                    # 主索引
                    term_map[key].extend(clean_values)
                    # 小写索引
                    term_map[key.lower()].extend(clean_values)
                    # 创建反向索引
                    for val in clean_values:
                        term_map[val].append(key)
            return term_map
        except Exception as e:
            raise RuntimeError(f"术语库加载失败：{str(e)}")

    @lru_cache(maxsize=3000)
    def _protected_translate(self, text: str) -> str:
        """精准回译引擎"""
        # 术语保护
        protected = text
        for term, marker in self.protected_terms.items():
            protected = protected.replace(term, marker)

        try:
            # 正向翻译
            en_text = self.translator(protected)[0]['translation_text']
            en_text = en_text.replace("singer", "vocal").replace("beat", "节奏")
            
            # 逆向翻译
            zh_text = self.back_translator(en_text)[0]['translation_text']
            
            # 术语还原
            for term, marker in self.protected_terms.items():
                zh_text = zh_text.replace(marker, term)
            
            # 错误校正
            correction_map = {
                "自动自控系统": "自动修音",
                "前奏": "声场定位",
                "粒子大小": "颗粒感"
            }
            for wrong, correct in correction_map.items():
                zh_text = zh_text.replace(wrong, correct)
            
            return zh_text
        except Exception as e:
            print(f"翻译异常: {str(e)}")
            return text

    def _dynamic_replace(self, text: str) -> str:
        """智能术语替换引擎"""
        words = jieba.lcut(text)
        cursor = 0
        replaced = []
        
        while cursor < len(words):
            found = False
            # 优先匹配长短语（3→2→1）
            for length in [3, 2, 1]:
                if cursor + length > len(words):
                    continue
                
                phrase = "".join(words[cursor:cursor+length])
                if replacements := self.term_map.get(phrase):
                    replaced.append(random.choice(replacements))
                    cursor += length
                    found = True
                    break
            if not found:
                # 单字替换（概率90%）
                current = words[cursor]
                if random.random() < 0.9:
                    replaced.append(random.choice(self.term_map.get(current, [current])))
                else:
                    replaced.append(current)
                cursor += 1
        return "".join(replaced)

    def process(self, text: str) -> dict:
        """生成增强结果"""
        return {
            "original": text,
            "backtrans": self._protected_translate(text),
            "eda1": self._dynamic_replace(text),
            "eda2": self._dynamic_replace(text)
        }

def main():
    # 初始化增强器
    try:
        augmentor = MusicAugmentor()
    except Exception as e:
        print(f"初始化失败: {str(e)}")
        return

    # 创建输出文档
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal'].font.size = Pt(11)

    try:
        input_doc = Document(INPUT_PATH)
        total = len(input_doc.paragraphs)
        
        for idx, para in enumerate(input_doc.paragraphs, 1):
            if text := para.text.strip():
                result = augmentor.process(text)
                doc.add_paragraph(f"原始数据：{result['original']}")
                doc.add_paragraph(f"回译增强：{result['backtrans']}")
                doc.add_paragraph(f"EDA变体1：{result['eda1']}")
                doc.add_paragraph(f"EDA变体2：{result['eda2']}\n")
                print(f"处理进度: {idx}/{total} ({idx/total:.1%})")

        doc.save(OUTPUT_PATH)
        print(f"\n✅ 处理完成！结果保存至: {OUTPUT_PATH}")

    except Exception as e:
        print(f"\n❌ 处理中断: {str(e)}")

if __name__ == "__main__":
    main()