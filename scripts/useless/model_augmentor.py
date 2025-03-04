"""
混音需求增强系统 - 模型增强版 v1.0
"""
import os
import json
import re
import time
import jieba
import torch
from docx import Document
from docx.shared import Pt
from transformers import AutoTokenizer, AutoModelForCausalLM
from concurrent.futures import ThreadPoolExecutor

# 配置路径
MODEL_PATH = "/Volumes/Study/prj/models/models--deepseek-ai--DeepSeek-R1-Distill-Qwen-1.5B/snapshots/530ca3e1ad39d440e182c2e4317aa40f012512fa"
INPUT_FILE = "/Volumes/Study/prj/data/raw/training_raw_data.docx"
OUTPUT_FILE = "/Volumes/Study/prj/data/processed/augmented_model.docx"

# 全局配置
MAX_TEXT_LENGTH = 150
MIN_TEXT_LENGTH = 8
BATCH_SIZE = 16

class ModelAugmentor:
    def __init__(self):
        self.device = torch.device("mps" if torch.backends.mps.is_available() else "cpu")
        self.tokenizer, self.model = self._load_model()
        self.executor = ThreadPoolExecutor(max_workers=2)
        jieba.initialize()
        print("✅ 模型初始化完成")

    def _load_model(self):
        try:
            tokenizer = AutoTokenizer.from_pretrained(MODEL_PATH, trust_remote_code=True)
            model = AutoModelForCausalLM.from_pretrained(
                MODEL_PATH,
                device_map=self.device,
                torch_dtype=torch.float16,
                trust_remote_code=True
            ).eval()
            print(f"🎛️ 模型加载成功 | 设备: {model.device}")
            return tokenizer, model
        except Exception as e:
            raise RuntimeError(f"模型加载失败: {str(e)}")

    def _generate_batch(self, texts):
        prompts = [f"生成专业混音变体：{text}" for text in texts]
        inputs = self.tokenizer(prompts, return_tensors="pt", padding=True, truncation=True, max_length=128).to(self.device)
        
        outputs = self.model.generate(
            inputs.input_ids,
            max_new_tokens=64,
            temperature=0.6,
            top_p=0.9,
            do_sample=True,
            pad_token_id=self.tokenizer.eos_token_id
        )
        
        return [
            re.findall(r"变体[:：]?\s*(.+?)(?=\n|$)", 
            self.tokenizer.decode(o, skip_special_tokens=True))[:3] 
            for o in outputs
        ]

    def process_text(self, text):
        variants = self._generate_batch([text])[0]
        return [
            v for v in variants
            if (MIN_TEXT_LENGTH <= len(v) <= MAX_TEXT_LENGTH)
            and not re.search(r"[\[\]【】()（）]", v)
        ][:3]

def main():
    augmentor = ModelAugmentor()
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    
    input_doc = Document(INPUT_FILE)
    paragraphs = [p.text.strip() for p in input_doc.paragraphs if p.text.strip()]
    
    start_time = time.time()
    for idx, text in enumerate(paragraphs, 1):
        try:
            variants = augmentor.process_text(text)
            doc.add_paragraph(f"原始需求：{text}").bold = True
            if variants:
                doc.add_paragraph("模型生成：" + " | ".join(variants))
            print(f"\r处理进度: {idx}/{len(paragraphs)} | 耗时: {time.time()-start_time:.1f}s", end="")
        except Exception as e:
            print(f"\n⚠️ 处理失败: {text[:20]}...")
    
    doc.save(OUTPUT_FILE)
    print(f"\n✅ 增强完成！保存至: {OUTPUT_FILE}")

if __name__ == "__main__":
    print("=== 模型增强系统 ===")
    start = time.time()
    try:
        main()
    except KeyboardInterrupt:
        print("\n操作已中断")
    finally:
        print(f"总耗时: {time.time()-start:.1f}秒")
        