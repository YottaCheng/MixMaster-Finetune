"""
混音需求增强系统 v7.2
(终极优化版)
"""
import os
import json
import re
import random
import time
import jieba
import torch
import numpy as np
from datetime import datetime
from packaging import version
from docx import Document
from docx.shared import Pt
from transformers import AutoTokenizer, AutoModelForCausalLM
from concurrent.futures import ThreadPoolExecutor

# 配置路径
CONFIG_PATH = "/Volumes/Study/prj/config/music_synonyms.json"
MODEL_PATH = "/Volumes/Study/prj/models/models--deepseek-ai--DeepSeek-R1-Distill-Qwen-1.5B/snapshots/530ca3e1ad39d440e182c2e4317aa40f012512fa"
INPUT_FILE = "/Volumes/Study/prj/data/raw/training_raw_data.docx"
OUTPUT_FILE = "/Volumes/Study/prj/data/processed/augmented_results_final.docx"

# 全局配置
MAX_TEXT_LENGTH = 150
MIN_TEXT_LENGTH = 8
BATCH_SIZE = 16  # 增大批量处理
CACHE_SIZE = 100  # 内存缓存条数

class HybridAugmentorPro:
    def __init__(self):
        # 硬件初始化
        self.device = torch.device("mps" if torch.backends.mps.is_available() else "cpu")
        self._configure_hardware()
        
        # 加载资源
        self.term_dict = self._load_terminology()
        self.tokenizer, self.model = self._load_model()
        self.executor = ThreadPoolExecutor(max_workers=2)
        jieba.initialize()
        print("✅ 系统初始化完成")

    def _configure_hardware(self):
        """硬件优化配置"""
        if self.device.type == "mps":
            torch.mps.set_per_process_memory_fraction(0.9)
            torch.mps.empty_cache()
            print("✅ MPS加速已启用")
            
    def _load_terminology(self):
        """加载术语库"""
        with open(CONFIG_PATH) as f:
            return {
                term: aliases if isinstance(aliases, list) else [aliases]
                for cat in json.load(f).values() for term, aliases in cat.items()
            }

    def _load_model(self):
        """优化模型加载"""
        try:
            tokenizer = AutoTokenizer.from_pretrained(MODEL_PATH, trust_remote_code=True)
            model = AutoModelForCausalLM.from_pretrained(
                MODEL_PATH,
                device_map=self.device,
                torch_dtype=torch.float16,
                trust_remote_code=True
            ).eval()
            print(f"🎛️ 模型加载成功 | 设备: {model.device}")
            return tokenizer, model
        except Exception as e:
            raise RuntimeError(f"模型加载失败: {str(e)}")

    def _term_replace(self, text):
        """增强型术语替换"""
        return ''.join([
            random.choice(self.term_dict.get(word, [word]))
            if random.random() < 0.85 else word 
            for word in jieba.lcut(text)
        ])

    def _generate_batch(self, texts):
        """批量生成优化"""
        prompts = [
            f"生成3个专业混音变体（仅术语）：{text}"
            for text in texts
        ]
        
        inputs = self.tokenizer(
            prompts, 
            return_tensors="pt", 
            padding=True, 
            truncation=True, 
            max_length=128
        ).to(self.device)
        
        outputs = self.model.generate(
            inputs.input_ids,
            max_new_tokens=64,
            temperature=0.5,
            top_p=0.85,
            do_sample=True,
            pad_token_id=self.tokenizer.eos_token_id
        )
        
        return [
            re.findall(r"变体[:：]?\s*(.+?)(?=\n|$)", 
            self.tokenizer.decode(o, skip_special_tokens=True))[:3] 
            for o in outputs
        ]

    def process_text(self, text):
        """处理流水线"""
        # 术语替换
        eda_variant = self._term_replace(text)
        
        # 模型生成
        model_variants = self._generate_batch([text])[0]
        
        # 结果过滤
        valid_variants = [
            v for v in [eda_variant] + model_variants
            if (8 <= len(v) <= 150) 
            and (v != text)
            and not re.search(r"[\[\]【】()（）]", v)
        ][:5]
        
        return valid_variants

def main():
    # 初始化系统
    augmentor = HybridAugmentorPro()
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    
    # 批量读取
    input_doc = Document(INPUT_FILE)
    paragraphs = [p.text.strip() for p in input_doc.paragraphs if p.text.strip()]
    
    # 缓存处理
    cache = []
    start_time = time.time()
    
    for idx, text in enumerate(paragraphs, 1):
        try:
            # 异步处理
            future = augmentor.executor.submit(augmentor.process_text, text)
            variants = future.result()
            
            # 写入缓存
            cache.append((text, variants))
            
            # 批量写入
            if len(cache) >= CACHE_SIZE or idx == len(paragraphs):
                for orig, vars in cache:
                    p = doc.add_paragraph()
                    p.add_run(f"原始需求：{orig}").bold = True
                    doc.add_paragraph("生成变体：" + " | ".join(vars))
                cache = []
                
            print(f"\r处理进度: {idx}/{len(paragraphs)} | 耗时: {time.time()-start_time:.1f}s", end="")
            
        except Exception as e:
            print(f"\n⚠️ 处理失败: {text[:20]}... | 错误: {str(e)}")

    # 保存结果
    doc.save(OUTPUT_FILE)
    print(f"\n✅ 处理完成！保存至: {OUTPUT_FILE}")

if __name__ == "__main__":
    print("=== 混音需求增强系统 v7.2 ===")
    start = time.time()
    try:
        main()
    except KeyboardInterrupt:
        print("\n🛑 用户中断操作")
    finally:
        print(f"总耗时: {time.time()-start:.1f}秒")