# music_augmentor.py
import os
import json
import random
import docx
from pathlib import Path
from transformers import pipeline, AutoTokenizer
import torch

class MusicAugmentor:
    def __init__(self):
        # 配置参数
        self.project_root = Path("/Volumes/Study/prj")
        self.synonym_path = self.project_root / "config/music_synonyms.json"
        self.raw_path = self.project_root / "data/raw/training_raw_data.docx"
        self.output_path = self.project_root / "data/augmented/enhanced_data.docx"
        
        # 加载资源
        self.synonyms = self._load_synonyms()
        self.generator = self._init_generator()
    
    def _load_synonyms(self):
        """加载领域同义词库"""
        with open(self.synonym_path, 'r') as f:
            return json.load(f)
    
    def _init_generator(self):
        """初始化适合M1/M3的轻量模型"""
        return pipeline(
            "text-generation",
            model="uer/gpt2-chinese-lyric",
            device="mps",
            torch_dtype=torch.float16,
            model_kwargs={
                "low_cpu_mem_usage": True,
                "pad_token_id": 50256  # 显式设置pad_token_id
            }
        )
    
    def _eda_replace(self, text):
        """专业领域同义词替换"""
        words = text.split()
        for i in range(len(words)):
            if words[i] in self.synonyms and random.random() < 0.6:
                words[i] = random.choice(self.synonyms[words[i]])
        return " ".join(words)
    
    def _generate_variants(self, text, num=5):
        """生成多样化变体"""
        prompt = f"用{num}种专业表达改写音乐制作需求，保持原意：{text}\n改写："
        
        try:
            outputs = self.generator(
                prompt,
                max_new_tokens=50,
                num_return_sequences=num,
                temperature=0.85,
                truncation=True,
                pad_token_id=50256  # 与模型设置一致
            )
            
            return [
                output['generated_text'].split("改写：")[-1]
                .strip()
                .replace("\n", "")
                for output in outputs
                if 10 < len(output['generated_text']) < 100
            ]
        except Exception as e:
            print(f"生成时出错: {str(e)[:50]}")
            return []

    def process(self):
        """处理主流程"""
        # 确保输出目录存在
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 读取原始数据
        doc = docx.Document(self.raw_path)
        raw_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # 执行增强
        enhanced = []
        for text in raw_texts[:50]:  # 先测试前50条
            try:
                # 保留原始数据
                enhanced.append(text)
                # EDA增强
                eda_text = self._eda_replace(text)
                if eda_text != text:
                    enhanced.append(eda_text)
                # 模型增强
                variants = self._generate_variants(text, num=3)
                enhanced.extend(variants)
            except Exception as e:
                print(f"处理失败: {text[:20]}... 错误: {str(e)[:50]}")
        
        # 去重并保存
        unique_data = list(set(enhanced))
        output_doc = docx.Document()
        
        # 添加带编号的内容
        for idx, text in enumerate(sorted(unique_data)[:400], 1):
            output_doc.add_paragraph(f"{idx}. {text}")
        
        # 确保文件可写
        if self.output_path.exists():
            self.output_path.unlink()
            
        output_doc.save(self.output_path)
        print(f"成功生成 {len(unique_data)} 条数据，已保存到 {self.output_path}")

if __name__ == "__main__":
    augmentor = MusicAugmentor()
    augmentor.process()