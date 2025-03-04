import os
import re
import math
import json
import random
from collections import defaultdict
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

class DataOptimizer:
    def __init__(self, input_path, output_path):
        self.input_path = input_path
        self.output_path = output_path
        self.data = []
        self.label_dist = defaultdict(int)
        self.max_entropy = math.log2(7)  # 7个有效标签的理论最大熵

    def load_data(self):
        """加载并预处理数据"""
        doc = Document(self.input_path)
        for para in doc.paragraphs:
            if text := para.text.strip():
                # 提取标签（支持多标签）
                parts = text.split('（')
                content = '（'.join(parts[:-1]).strip()
                labels = [l.strip() for l in parts[-1].rstrip('）').split('、')]
                self.data.append({
                    "text": content,
                    "labels": labels
                })
        print(f"加载{len(self.data)}条原始数据")

    def calculate_entropy(self):
        """计算多标签信息熵"""
        total = len(self.data)
        label_counts = defaultdict(int)
        for item in self.data:
            for label in item['labels']:
                label_counts[label] += 1
        
        entropy = 0.0
        for count in label_counts.values():
            prob = count / total
            entropy -= prob * math.log2(prob) if prob > 0 else 0
        return entropy, label_counts

    def deduplicate(self):
        """严格去重策略"""
        seen = set()
        unique_data = []
        for item in self.data:
            key = (item['text'], tuple(sorted(item['labels'])))
            if key not in seen:
                seen.add(key)
                unique_data.append(item)
        self.data = unique_data
        print(f"去重后剩余{len(self.data)}条数据")

    def filter_similar_samples(self, threshold=0.8):
        """基于TF-IDF的相似度过滤"""
        texts = [item['text'] for item in self.data]
        vectorizer = TfidfVectorizer()
        tfidf = vectorizer.fit_transform(texts)
        
        similarity_matrix = cosine_similarity(tfidf)
        to_remove = set()
        
        for i in range(len(similarity_matrix)):
            for j in range(i+1, len(similarity_matrix)):
                if similarity_matrix[i,j] > threshold:
                    to_remove.add(j)
        
        filtered_data = [item for idx, item in enumerate(self.data) 
                        if idx not in to_remove]
        self.data = filtered_data
        print(f"相似度过滤后剩余{len(self.data)}条数据")

    def adjust_distribution(self, target_entropy=2.65, max_iterations=50):
        """动态调整标签分布"""
        current_entropy, label_dist = self.calculate_entropy()
        print(f"当前熵值: {current_entropy:.4f} (目标: {target_entropy})")
        
        iteration = 0
        while current_entropy > target_entropy and iteration < max_iterations:
            iteration += 1
            print(f"第{iteration}次调整...")
            
            # 找出超出平均分布的标签
            avg = len(self.data) / len(label_dist)
            oversampled = [label for label, count in label_dist.items() 
                          if count > avg * 1.5]
            
            # 如果没有明显过多样本，则退出
            if not oversampled:
                print("未发现明显过多样本，停止调整")
                break
            
            # 随机删除高频标签样本
            label = max(oversampled, key=lambda x: label_dist[x])
            candidates = [item for item in self.data if label in item['labels']]
            remove_count = max(1, int(len(candidates)*0.2))  # 增加删除比例到20%
            self.data = [item for item in self.data 
                        if item not in random.sample(candidates, remove_count)]
            
            # 更新熵值和标签分布
            current_entropy, label_dist = self.calculate_entropy()
            print(f"删除{remove_count}条'{label}'样本，当前熵值: {current_entropy:.4f}")
        
        if iteration >= max_iterations:
            print("达到最大迭代次数，停止调整")

    def save_results(self):
        """保存优化后的数据"""
        doc = Document()
        for idx, item in enumerate(self.data, 1):
            p = doc.add_paragraph(f"{idx}. {item['text']}")
            p.add_run(f"（{'、'.join(item['labels'])}）").bold = True
        doc.save(self.output_path)
        print(f"优化数据已保存至：{self.output_path}")

    def generate_report(self):
        """生成数据质量报告"""
        entropy, label_dist = self.calculate_entropy()
        report = {
            "total_samples": len(self.data),
            "unique_labels": len(label_dist),
            "entropy": round(entropy, 4),
            "label_distribution": dict(label_dist),
            "quality_metrics": {
                "duplicate_rate": 1 - (len(self.data)/original_count),
                "similarity_threshold": 0.8,
                "adjustments_made": True if entropy < self.max_entropy else False
            }
        }
        with open("optimization_report.json", "w") as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        print("优化报告已生成")

if __name__ == "__main__":
    optimizer = DataOptimizer(
        input_path="/Volumes/Study/prj/data/raw/training_labeled_data.docx",
        output_path="/Volumes/Study/prj/data/processed/optimized_data.docx"
    )
    
    original_count = 772  # 初始样本数
    optimizer.load_data()
    optimizer.deduplicate()
    optimizer.filter_similar_samples()
    optimizer.adjust_distribution(target_entropy=2.65)
    optimizer.generate_report()
    optimizer.save_results()